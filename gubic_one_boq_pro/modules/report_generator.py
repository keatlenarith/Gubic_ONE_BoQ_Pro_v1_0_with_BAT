"""Report export helpers for Excel, DOCX, and PDF."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd

from config.settings import EXPORT_DIR
from modules.cost_engine import calculate_kpis, cost_breakdown, package_ranking, top_items, validation_report
from modules.material_engine import material_summary
from modules.labor_engine import labor_summary
from modules.equipment_engine import equipment_summary


def _stamp(name: str, suffix: str) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in name)[:70]
    return EXPORT_DIR / f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{suffix}"


def export_detailed_excel(df: pd.DataFrame, project_name: str = "BoQ_Project") -> Path:
    path = _stamp(f"{project_name}_Detailed_BoQ", "xlsx")
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Standardized_BoQ", index=False)
        pd.DataFrame([calculate_kpis(df)]).to_excel(writer, sheet_name="KPI_Summary", index=False)
        cost_breakdown(df).to_excel(writer, sheet_name="Cost_Breakdown", index=False)
        package_ranking(df).to_excel(writer, sheet_name="Package_Ranking", index=False)
        top_items(df, 20).to_excel(writer, sheet_name="Top_20_Items", index=False)
        validation_report(df).to_excel(writer, sheet_name="Validation_Report", index=False)
    return path


def export_word_report(df: pd.DataFrame, project_name: str = "BoQ Project") -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    path = _stamp(f"{project_name}_BoQ_Summary_Report", "docx")
    kpis = calculate_kpis(df)
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("Gubic ONE BoQ Pro - Executive BoQ Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}")

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"The imported BoQ contains {kpis['number_of_boq_items']:,} standardized item rows across "
        f"{kpis['number_of_packages']:,} packages/sheets. The total project cost is "
        f"${kpis['total_project_cost']:,.2f}. Based on an area of {kpis['area_m2']:,.2f} m², "
        f"the calculated unit cost is ${kpis['cost_per_m2']:,.2f}/m²."
    )

    doc.add_heading("2. Cost Summary", level=1)
    summary = doc.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    summary.rows[0].cells[0].text = "Metric"
    summary.rows[0].cells[1].text = "Value"
    for label, value in [
        ("Total Project Cost", f"${kpis['total_project_cost']:,.2f}"),
        ("Material Cost", f"${kpis['total_material_cost']:,.2f}"),
        ("Labor Cost", f"${kpis['total_labor_cost']:,.2f}"),
        ("Equipment Cost", f"${kpis['total_equipment_cost']:,.2f}"),
        ("Transport Cost", f"${kpis['total_transport_cost']:,.2f}"),
        ("Risk / Contingency", f"${kpis['total_risk_cost']:,.2f}"),
        ("Cost per m²", f"${kpis['cost_per_m2']:,.2f}/m²"),
    ]:
        row = summary.add_row().cells
        row[0].text = label
        row[1].text = value

    def add_df_table(title: str, table_df: pd.DataFrame, max_rows: int = 10):
        doc.add_heading(title, level=1)
        if table_df.empty:
            doc.add_paragraph("No data available.")
            return
        table_df = table_df.head(max_rows).copy()
        table = doc.add_table(rows=1, cols=len(table_df.columns))
        table.style = "Table Grid"
        for i, c in enumerate(table_df.columns):
            table.rows[0].cells[i].text = str(c)
        for _, rec in table_df.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(table_df.columns):
                val = rec[c]
                if isinstance(val, float):
                    cells[i].text = f"{val:,.2f}"
                else:
                    cells[i].text = str(val)

    add_df_table("3. Cost Breakdown", cost_breakdown(df), 10)
    add_df_table("4. Top 20 Expensive Items", top_items(df, 20), 20)
    add_df_table("5. Material Summary", material_summary(df), 10)
    add_df_table("6. Labor Summary", labor_summary(df), 10)
    add_df_table("7. Equipment Summary", equipment_summary(df), 10)

    doc.add_heading("8. Key Observations", level=1)
    breakdown = cost_breakdown(df)
    if not breakdown.empty:
        dominant = breakdown.iloc[0]
        doc.add_paragraph(f"The dominant cost type is {dominant['cost_type']} at ${dominant['amount']:,.2f}.")
    rank = package_ranking(df).head(1)
    if not rank.empty:
        doc.add_paragraph(f"The highest-cost package/sheet is {rank.iloc[0]['source_sheet']} at ${rank.iloc[0]['total_cost']:,.2f}.")
    doc.add_paragraph("Review the validation report before issuing the BoQ commercially, especially amount mismatches, missing quantities/rates, and duplicate item codes.")

    doc.save(path)
    return path


def export_pdf_summary(df: pd.DataFrame, project_name: str = "BoQ Project") -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    path = _stamp(f"{project_name}_Executive_BoQ_Summary", "pdf")
    kpis = calculate_kpis(df)
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story: list[Any] = []
    story.append(Paragraph("Gubic ONE BoQ Pro - Executive BoQ Summary", styles["Title"]))
    story.append(Paragraph(f"Project: {project_name}", styles["Normal"]))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 12))
    data = [
        ["Metric", "Value"],
        ["Total Project Cost", f"${kpis['total_project_cost']:,.2f}"],
        ["Direct Cost", f"${kpis['total_direct_cost']:,.2f}"],
        ["Material Cost", f"${kpis['total_material_cost']:,.2f}"],
        ["Labor Cost", f"${kpis['total_labor_cost']:,.2f}"],
        ["Cost per m²", f"${kpis['cost_per_m2']:,.2f}/m²"],
        ["BoQ Items", f"{kpis['number_of_boq_items']:,}"],
    ]
    table = Table(data, colWidths=[220, 220])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B365D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9E2EC")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FC")]),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph("Top Cost Items", styles["Heading2"]))
    top = top_items(df, 10)
    if not top.empty:
        tdata = [["Sheet", "Description", "Cost"]]
        for _, r in top.iterrows():
            tdata.append([str(r.get("source_sheet", ""))[:22], str(r.get("item_description", ""))[:52], f"${r.get('total_cost', 0):,.2f}"])
        t = Table(tdata, colWidths=[105, 285, 90])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B365D")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9E2EC")),
        ]))
        story.append(t)
    doc.build(story)
    return path
